def inv():
     from docx import Documents
     from docx.enum.style import WD_STYLE_TYPE
     from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
     from docx.shared import Pt

     Intro= "It would be a pleasure to have the company of"
     Address="at 11010 Memory Lane on the Evening of"
     Date="April 1st"
     Time="at 7 o'clock"

     with open('./guests.txt') as file:
         guestList= file.read().splitlines()

     document = Documents()
     styles = document.styles

     style = styles.add_style("start", WD_STYLE_TYPE.PARAGRAPH)
     font = style.font
     font.name = "MathJax_Caligraphic"  
     font.size = Pt(12)
     font.all_caps = True  
     font.bold = True
     style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

     style = styles.add_style("date", WD_STYLE_TYPE.PARAGRAPH)
     font = style.font
     font.name = "Liberation Sans"
     font.size = Pt(12)
     font.bold= False
     style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

     style = styles.add_style("name", WD_STYLE_TYPE.PARAGRAPH)
     font = style.font
     font.name = "Liberation Sans"  
     font.size = Pt(16)
     font.bold = True
     style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



     for guests in guestList.txt:
        document.add_paragraph(Intro, style="start")
        document.add_paragraph(guests, style="name")
        document.add_paragraph(Address, style="start")
        document.add_paragraph(Date, style="date")
        document.add_paragraph(Time, style="start")

     document.save('Invitations.docx')
if __name__=='__inv__':
    inv()